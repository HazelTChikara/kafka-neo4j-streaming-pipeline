"""
Generates a 4-page IEEE-style project proposal Word document (.docx) for:
Directional Stock Price Prediction Using News Data
Authors: Hazel Chikara, Tichaona Dzitiro, Praise Makore,
         Tinotenda Muponda, Takudzwa Taanisa, Tinashe Tomu
"""

from docx import Document
from docx.shared import Pt, Inches, RGBColor, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
import os

OUTPUT_PATH = os.path.join(os.path.dirname(__file__), "Project_Proposal.docx")

doc = Document()

# ── IEEE page margins ──────────────────────────────────────────────────────────
for section in doc.sections:
    section.top_margin    = Cm(1.9)
    section.bottom_margin = Cm(2.0)
    section.left_margin   = Cm(1.7)
    section.right_margin  = Cm(1.7)

BODY_FONT = "Times New Roman"
BODY_SIZE = 10
ABST_SIZE = 9

# ── Helpers ────────────────────────────────────────────────────────────────────
def set_run(run, name=BODY_FONT, size=BODY_SIZE, bold=False, italic=False, color=None):
    run.font.name   = name
    run.font.size   = Pt(size)
    run.font.bold   = bold
    run.font.italic = italic
    if color:
        run.font.color.rgb = color

def para_fmt(p, before=0, after=4, align=None, left=0, first=0):
    pf = p.paragraph_format
    pf.space_before = Pt(before)
    pf.space_after  = Pt(after)
    if left:
        pf.left_indent = Pt(left)
    if first:
        pf.first_line_indent = Pt(first)
    if align is not None:
        p.alignment = align

def body_para(text, bold=False, italic=False, before=0, after=4,
              align=WD_ALIGN_PARAGRAPH.JUSTIFY, first_indent=14):
    p = doc.add_paragraph()
    para_fmt(p, before=before, after=after, align=align, first=first_indent)
    r = p.add_run(text)
    set_run(r, bold=bold, italic=italic)
    return p

def section_heading(text, number=None, before=8, after=3):
    p = doc.add_paragraph()
    para_fmt(p, before=before, after=after, align=WD_ALIGN_PARAGRAPH.CENTER)
    label = f"{number}. {text.upper()}" if number else text.upper()
    r = p.add_run(label)
    set_run(r, size=BODY_SIZE, bold=True)
    return p

def subsection_heading(text, label, before=5, after=1):
    p = doc.add_paragraph()
    para_fmt(p, before=before, after=after, align=WD_ALIGN_PARAGRAPH.LEFT)
    r = p.add_run(f"{label} {text}")
    set_run(r, size=BODY_SIZE, bold=True, italic=True)
    return p

def add_rule(before=2, after=2):
    p = doc.add_paragraph()
    para_fmt(p, before=before, after=after)
    pPr  = p._p.get_or_add_pPr()
    pBdr = OxmlElement("w:pBdr")
    bot  = OxmlElement("w:bottom")
    bot.set(qn("w:val"),   "single")
    bot.set(qn("w:sz"),    "4")
    bot.set(qn("w:space"), "1")
    bot.set(qn("w:color"), "000000")
    pBdr.append(bot)
    pPr.append(pBdr)

def shd_cell(cell, hex_fill):
    tc   = cell._tc
    tcPr = tc.get_or_add_tcPr()
    shd  = OxmlElement("w:shd")
    shd.set(qn("w:val"),   "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"),  hex_fill)
    tcPr.append(shd)

def add_table(doc, headers, rows, col_widths=None):
    table = doc.add_table(rows=1 + len(rows), cols=len(headers))
    table.style     = "Table Grid"
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    hcells = table.rows[0].cells
    for i, h in enumerate(headers):
        hcells[i].text = h
        for run in hcells[i].paragraphs[0].runs:
            set_run(run, size=ABST_SIZE, bold=True, color=RGBColor(0xFF, 0xFF, 0xFF))
        hcells[i].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
        shd_cell(hcells[i], "1F3864")
    for ri, row_data in enumerate(rows):
        rc = table.rows[ri + 1].cells
        for ci, val in enumerate(row_data):
            rc[ci].text = val
            for run in rc[ci].paragraphs[0].runs:
                set_run(run, size=ABST_SIZE)
            if ri % 2 == 1:
                shd_cell(rc[ci], "DCE6F1")
    if col_widths:
        for i, w in enumerate(col_widths):
            for row in table.rows:
                row.cells[i].width = Inches(w)
    sp = doc.add_paragraph()
    para_fmt(sp, before=0, after=2)
    return table

# ══════════════════════════════════════════════════════════════════════════════
# TITLE BLOCK
# ══════════════════════════════════════════════════════════════════════════════
title_p = doc.add_paragraph()
para_fmt(title_p, before=0, after=6, align=WD_ALIGN_PARAGRAPH.CENTER)
t_run = title_p.add_run("Directional Stock Price Prediction Using News Data")
set_run(t_run, size=20, bold=False)

# Authors in a borderless 6-column table — names only, side by side
AUTHORS = [
    "Hazel Chikara",
    "Tichaona Dzitiro",
    "Praise Makore",
    "Tinotenda Muponda",
    "Takudzwa Taanisa",
    "Tinashe Tomu",
]
auth_tbl = doc.add_table(rows=1, cols=len(AUTHORS))
auth_tbl.style     = "Table Grid"
auth_tbl.alignment = WD_TABLE_ALIGNMENT.CENTER

for i, name in enumerate(AUTHORS):
    cell = auth_tbl.rows[0].cells[i]
    # Remove all borders
    tc   = cell._tc
    tcPr = tc.get_or_add_tcPr()
    tcBdr = OxmlElement("w:tcBdr")
    for side in ("top", "left", "bottom", "right", "insideH", "insideV"):
        b = OxmlElement(f"w:{side}")
        b.set(qn("w:val"),   "none")
        b.set(qn("w:sz"),    "0")
        b.set(qn("w:space"), "0")
        b.set(qn("w:color"), "auto")
        tcBdr.append(b)
    tcPr.append(tcBdr)
    # Write name only
    cp = cell.paragraphs[0]
    cp.alignment = WD_ALIGN_PARAGRAPH.CENTER
    para_fmt(cp, before=2, after=2)
    cr = cp.add_run(name)
    set_run(cr, size=10, bold=False)

sp = doc.add_paragraph()
para_fmt(sp, before=4, after=0)
add_rule(before=2, after=2)

# ══════════════════════════════════════════════════════════════════════════════
# ABSTRACT
# ══════════════════════════════════════════════════════════════════════════════
abst_p = doc.add_paragraph()
para_fmt(abst_p, before=2, after=0, align=WD_ALIGN_PARAGRAPH.JUSTIFY, left=36, first=0)
set_run(abst_p.add_run("Abstract\u2014"), size=ABST_SIZE, bold=True)
set_run(abst_p.add_run(
    "The fusion of financial news sentiment with quantitative market data offers a compelling "
    "frontier for improving stock market forecasting [1]. This proposal describes the creation "
    "and assessment of machine learning models for forecasting directional stock price movement. "
    "Using a dataset of news articles and historical stock charts for Apple (AAPL) and Amazon "
    "(AMZN), we compare classical classifiers including Support Vector Machines (SVM) and "
    "Logistic Regression, the high-performance ensemble method XGBoost [2], and the "
    "state-of-the-art transformer FinBERT for sophisticated sentiment feature extraction [3]. "
    "The primary objective is to determine which modeling approach most effectively deciphers "
    "the complex interplay between public sentiment and stock price dynamics. Model performance "
    "will be rigorously assessed using 10-fold cross-validation and key classification metrics "
    "including F1-score, precision, and recall. This research provides empirical insights into "
    "the predictive power of integrating modern NLP features with established machine learning "
    "techniques for financial forecasting."
), size=ABST_SIZE)

kw_p = doc.add_paragraph()
para_fmt(kw_p, before=3, after=6, align=WD_ALIGN_PARAGRAPH.JUSTIFY, left=36, first=0)
set_run(kw_p.add_run("Index Terms\u2014"), size=ABST_SIZE, bold=True, italic=True)
set_run(kw_p.add_run(
    "Stock Price Prediction, Sentiment Analysis, Machine Learning, "
    "Natural Language Processing, FinBERT, Financial Forecasting."
), size=ABST_SIZE, italic=True)

add_rule(before=2, after=4)

# ══════════════════════════════════════════════════════════════════════════════
# I. INTRODUCTION
# ══════════════════════════════════════════════════════════════════════════════
section_heading("Introduction", "I")
body_para(
    "Predicting the direction of stock prices remains a formidable challenge in financial "
    "markets, with critical ramifications for investment strategy and risk management [5]. "
    "For decades, analysts have relied on traditional methods like technical and fundamental "
    "analysis. However, these approaches, primarily based on historical price data and company "
    "financials, often fail to capture the real-time, sentiment-driven volatility of modern "
    "markets."
)
body_para(
    "The digital revolution has introduced a paradigm shift, unlocking massive unstructured data "
    "streams from financial news, social media, and analyst reports [4]. This textual data "
    "provides a powerful real-time lens into market sentiment—the collective psychology of "
    "investors—which is a crucial but often overlooked driver of price movements. Recent research "
    "increasingly confirms that such data contains useful predictive signals."
)
body_para(
    "This project harnesses these signals using advanced machine learning and NLP techniques. "
    "We develop and validate a suite of predictive models by combining sentiment from news "
    "articles with conventional time-series stock data, implementing models from classical ML "
    "to the state-of-the-art, domain-specific transformer FinBERT, providing a clear, "
    "evidence-based comparative analysis."
)

# ══════════════════════════════════════════════════════════════════════════════
# II. PROBLEM DEFINITION
# ══════════════════════════════════════════════════════════════════════════════
section_heading("Problem Definition", "II")
body_para(
    "The central objective is to forecast the directional movement of a stock's price as a "
    "binary classification problem: the model predicts whether the price will move \u201cup\u201d "
    "or \u201cdown\u201d over a specified period. This directly challenges the Efficient Market "
    "Hypothesis (EMH), which holds that asset prices follow a random walk. Nonetheless, "
    "an increasing body of research suggests that news and social media sentiment can reveal "
    "predictable patterns and market inefficiencies."
)
body_para(
    "Our hypothesis is that by methodically measuring the sentiment in financial news, we can "
    "identify predictive features that, paired with historical price data, enable a machine "
    "learning model to predict directional movements with accuracy noticeably higher than "
    "chance. The framework is evaluated on Apple (AAPL) and Amazon (AMZN) stocks."
)

# Formula
fp = doc.add_paragraph()
para_fmt(fp, before=3, after=3, align=WD_ALIGN_PARAGRAPH.CENTER)
set_run(fp.add_run("y(t) = 1  if  Close(t+1) > Close(t),   else  y(t) = 0"),
        name="Courier New", size=10, bold=True)

# ══════════════════════════════════════════════════════════════════════════════
# III. DATASETS
# ══════════════════════════════════════════════════════════════════════════════
section_heading("Datasets", "III")
body_para(
    "This study\u2019s predictive framework is built on a multi-modal dataset integrating "
    "three sources: financial news articles, historical stock trading records, and real-time "
    "social media activity, focusing on Apple (AAPL) and Amazon (AMZN)."
)

subsection_heading("News Data", "A.")
body_para(
    "A corpus of financial news articles from NASDAQ published between December 2017 and "
    "February 2019 comprises roughly 78,000\u201387,000 articles stored in JSON format with "
    "headline, content, and publication date. FinBERT extracts sentiment scores (positive, "
    "negative, neutral) or rich contextual embeddings to capture each article\u2019s underlying "
    "market tone.", first_indent=0
)

subsection_heading("Stock Market Data", "B.")
body_para(
    "Historical daily OHLCV data for Apple and Amazon is obtained from Yahoo Finance. The "
    "binary target variable is derived from closing price changes: Up (1) for a price increase, "
    "Down (0) otherwise. Additional numerical features include trading volume and technical "
    "indicators such as volatility and moving averages.", first_indent=0
)

subsection_heading("Social Media Data (Twitter)", "C.")
body_para(
    "Twitter data related to Apple and Amazon is collected to capture real-time public "
    "sentiment. While formal news provides curated analysis, Twitter reflects speculative "
    "buzz and herd behavior that may precede market shifts. Tweets are processed to extract "
    "sentiment scores that, combined with news sentiment and market data, provide a holistic "
    "view of stock price drivers.", first_indent=0
)

# ══════════════════════════════════════════════════════════════════════════════
# IV. RESEARCH PLAN
# ══════════════════════════════════════════════════════════════════════════════
section_heading("Research Plan", "IV")

subsection_heading("Data Preprocessing", "A.")
body_para(
    "Raw sources include financial news, social media, and stock chart data for Apple and "
    "Amazon. Articles are separated into per-company CSV files containing title, text, "
    "timestamp, and social media share counts. All timestamps are standardized to GMT; "
    "date and time are stored in separate columns for alignment with stock data. Social "
    "media reach (Facebook, LinkedIn, Pinterest, Google+) is aggregated into a single "
    "scalar. For Yahoo Finance stock CSVs, descriptive headers are added, formats "
    "standardized, duplicates removed, and data sorted chronologically.", first_indent=0
)

subsection_heading("Feature Engineering", "B.")
body_para(
    "Text preprocessing encompasses tokenization, stopword removal, lemmatization, and "
    "normalization. N-gram modeling (unigrams, bigrams, trigrams) captures local word "
    "dependencies conveying market sentiment. Financial sentiment analysis is performed "
    "on news and Twitter data using transformer architectures (BERT/FinBERT) to derive "
    "polarity scores timestamped for alignment with price movements.", first_indent=0
)

# ══════════════════════════════════════════════════════════════════════════════
# V. MODELS AND ALGORITHMS
# ══════════════════════════════════════════════════════════════════════════════
section_heading("Models and Algorithms", "V")
body_para(
    "Four supervised learning algorithms are implemented: Logistic Regression (LR), "
    "Support Vector Machine (SVM), Extreme Gradient Boosting (XGBoost), and FinBERT, "
    "to predict the directional movement of AAPL and AMZN stock prices."
)

subsection_heading("Logistic Regression (LR)", "A.")
body_para(
    "LR serves as the interpretable binary classification baseline, estimating the "
    "probability of upward or downward price movement from sentiment polarity scores, "
    "tweet volume, and technical features. L2 regularization reduces overfitting, and "
    "model coefficients reveal the strongest correlates of price direction [6],[7].", first_indent=0
)

subsection_heading("Support Vector Machine (SVM)", "B.")
body_para(
    "SVM identifies an optimal hyperplane separating positive and negative market "
    "directions by maximizing the decision margin. Both linear and RBF kernels are "
    "evaluated to capture potential non-linear relationships between linguistic sentiment "
    "and price movements. SVM consistently outperforms traditional classifiers in "
    "sentiment-based forecasting [8],[9].", first_indent=0
)

subsection_heading("Extreme Gradient Boosting (XGBoost)", "C.")
body_para(
    "XGBoost builds multiple weak learners through gradient boosting to capture complex "
    "non-linear dependencies between sentiment features and technical indicators. Tree "
    "pruning, L1/L2 regularization, and subsampling prevent overfitting while maintaining "
    "efficiency. XGBoost is widely recognized as highly effective for multi-modal stock "
    "movement prediction [10],[11].", first_indent=0
)

subsection_heading("FinBERT", "D.")
body_para(
    "FinBERT is pre-trained on large-scale financial text and employs the BERT architecture "
    "to analyze contextual and directional dependencies. It converts financial news and "
    "tweets into dense contextual embeddings fine-tuned for directional sentiment "
    "classification, interpreting subtle tone variations and financial terminology [12],[13].", first_indent=0
)

subsection_heading("Training and Evaluation", "E.")
body_para(
    "All models are trained on chronologically ordered data using an 80\u201320 split and "
    "evaluated through 10-fold cross-validation. Performance metrics include Accuracy, "
    "Precision, Recall, and F1-Score to determine the most effective approach for "
    "integrating textual sentiment features and numerical market signals.", first_indent=0
)

# ══════════════════════════════════════════════════════════════════════════════
# VI. EVALUATION PLAN
# ══════════════════════════════════════════════════════════════════════════════
section_heading("Evaluation Plan", "VI")
body_para(
    "We evaluate Logistic Regression as a baseline, SVM for high-dimensional feature spaces, "
    "XGBoost as an effective gradient boosting method, and FinBERT for contextual understanding "
    "of financial text. The dataset is aligned to stock price movements, partitioned with an "
    "80\u201320 training\u2013testing split, and subjected to 10-fold cross-validation."
)
body_para(
    "Performance is quantified using F1-score (balancing precision and recall), individual "
    "precision and recall scores, and overall accuracy. Results will be compared against "
    "established benchmarks in directional stock price prediction research, providing a "
    "clear, evidence-based comparative analysis."
)

# ══════════════════════════════════════════════════════════════════════════════
# VII. PROJECT TIMELINE AND DIVISION OF WORK
# ══════════════════════════════════════════════════════════════════════════════
section_heading("Project Timeline and Division of Work", "VII")

tbl_lbl = doc.add_paragraph()
para_fmt(tbl_lbl, before=2, after=2, align=WD_ALIGN_PARAGRAPH.CENTER)
set_run(tbl_lbl.add_run(
    "TABLE I\nPROJECT PLAN WITH TASKS, DESCRIPTIONS, DEADLINES, AND RESPONSIBILITIES"
), size=ABST_SIZE, bold=True)

add_table(doc,
    headers=["Task", "Description", "Deadline", "Members"],
    rows=[
        ["Literature Review",
         "Review and analyze prior papers to refine data, establish performance benchmarks, and finalize methodology.",
         "10/10/2025", "All Members"],
        ["Data Cleaning",
         "Preprocess raw text; remove irrelevant characters, stopwords; standardize format; handle missing stock data.",
         "10/17/2025", "Hazel, Tichaona"],
        ["Feature Engineering",
         "Create numerical features from text using TF-IDF / sentiment lexicons; align news data with stock labels.",
         "10/24/2025", "Praise, Tinashe"],
        ["Baseline Model",
         "Implement classical ML models: Logistic Regression and SVM.",
         "10/31/2025", "Tinotenda, Takudzwa"],
        ["Advanced Model",
         "Implement XGBoost and FinBERT transformer for deep learning analysis.",
         "11/07/2025", "Tinashe, Takudzwa"],
        ["Model Training",
         "Train all four models using 10-fold cross-validation for robustness.",
         "11/11/2025", "Tinotenda, Tinashe, Takudzwa"],
        ["Presentation",
         "Create slide deck summarizing key findings; rehearse for final presentation.",
         "11/11/2025", "All Members"],
        ["Hyperparameter Tuning",
         "Systematically optimize parameters for each model to maximize predictive performance.",
         "11/11/2025", "All Members"],
        ["Performance Evaluation",
         "Run tuned models on held-out test set; compute F1-score, precision, and recall.",
         "11/28/2025", "Hazel, Tinotenda"],
        ["Comparative Analysis",
         "Analyze and visualize results to compare performance of the four models.",
         "12/02/2025", "Tichaona, Praise"],
        ["Final Report",
         "Draft the final academic report detailing methodology, results, analysis, and conclusions.",
         "12/07/2025", "All Members"],
    ],
    col_widths=[1.2, 3.2, 0.9, 1.0]
)

# ══════════════════════════════════════════════════════════════════════════════
# REFERENCES
# ══════════════════════════════════════════════════════════════════════════════
section_heading("References", before=6, after=2)

REFS = [
    "Shobayo, O.; Adeyemi-Longe, S.; Popoola, O.; Ogunleye, B. Innovative Sentiment Analysis and Prediction of Stock Price Using FinBERT, GPT-4 and Logistic Regression. Big Data Cogn. Comput. 2024, 8, 143. https://doi.org/10.3390/bdcc8110143",
    "J. Clerk Maxwell, A Treatise on Electricity and Magnetism, 3rd ed., vol. 2. Oxford: Clarendon, 1892, pp. 68\u201373.",
    "\u201cDirectional stock market forecast using deep learning models,\u201d GitHub repository, lee14257/stock-market-forecast-fyp. https://github.com/lee14257/stock-market-forecast-fyp (accessed Oct. 8, 2025).",
    "Du S, Shen H. A Stock Prediction Method Based on Deep Reinforcement Learning and Sentiment Analysis. Applied Sciences. 2024; 14(19):8747. https://doi.org/10.3390/app14198747",
    "A. Babu, \u201cPredicting Stock Prices Using LSTMs: A Step-by-Step Guide to Time Series Forecasting,\u201d Medium, Mar. 21, 2025. https://medium.com/@aditib259/predicting-stock-prices-using-lstms (accessed Oct. 8, 2025).",
    "M. Ballings, D. Van den Poel, N. Hespeels, and R. Gryp, \u201cEvaluating multiple classifiers for stock price direction prediction,\u201d Expert Syst. Appl., vol. 42, no. 20, pp. 7046\u20137056, 2015.",
    "A. K. Nassirtoussi, S. Aghabozorgi, T. Y. Wah, and D. C. L. Ngo, \u201cText mining of news headlines for forex market prediction,\u201d Expert Syst. Appl., vol. 42, no. 1, pp. 306\u2013324, 2015.",
    "J. Xiao et al., \u201cA new approach for stock price analysis and prediction based on SSA and SVM,\u201d Int. J. Inf. Technol. Decis. Mak., vol. 18, no. 1, pp. 287\u2013310, 2019.",
    "V. S. Pagolu, K. N. Reddy, G. Panda, and B. Majhi, \u201cSentiment Analysis of Twitter Data for Predicting Stock Market Movements,\u201d in Proc. SCOPES 2016, pp. 1345\u20131350.",
    "S. Basak et al., \u201cPredicting the direction of stock market prices using tree-based classifiers,\u201d N. Am. J. Econ. Finance, vol. 47, pp. 552\u2013567, 2019.",
    "B. Weng, M. A. Ahmed, and F. M. Megahed, \u201cStock market one-day ahead movement prediction using disparate data sources,\u201d Expert Syst. Appl., vol. 79, pp. 153\u2013163, 2017.",
    "Z. Liu, D. Huang, K. Huang, Z. Li, and J. Zhao, \u201cFinBERT: A Pre-Trained Financial Language Representation Model for Financial Text Mining,\u201d in Proc. IJCAI 2021, pp. 4513\u20134519.",
    "F. Yang, X. Sun, and J. Li, \u201cHybrid transformer-based framework for multi-source news-driven stock movement prediction,\u201d Expert Syst. Appl., vol. 219, pp. 119\u2013134, 2023.",
]

for i, ref in enumerate(REFS, 1):
    rp = doc.add_paragraph()
    para_fmt(rp, before=1, after=1, align=WD_ALIGN_PARAGRAPH.JUSTIFY, left=18, first=-18)
    set_run(rp.add_run(f"[{i}] {ref}"), size=ABST_SIZE)

# ── Save ───────────────────────────────────────────────────────────────────────
doc.save(OUTPUT_PATH)
print(f"\u2705  Proposal saved to: {OUTPUT_PATH}")
